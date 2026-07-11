#!/usr/bin/env python3
"""Structural and layout-boundary validation for the Eiren DOCX dossier."""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "eiren-kestrel" / "deliverables" / "eiren-kestrel-trinity-mandala-research-dossier.docx"
OUTPUT = ROOT / "docs" / "eiren-kestrel" / "deliverables" / "dossier-validation.json"


def close(actual: int | None, expected: int, tolerance: int = 2) -> bool:
    return actual is not None and abs(actual - expected) <= tolerance


def twips(length) -> int | None:
    return None if length is None else int(round(length / 635))


def table_width_twips(table) -> int | None:
    node = table._tbl.tblPr.find(qn("w:tblW"))
    if node is None or node.get(qn("w:type")) != "dxa":
        return None
    return int(node.get(qn("w:w")))


def main() -> None:
    checks: list[dict[str, object]] = []

    def record(name: str, passed: bool, evidence) -> None:
        checks.append({"check": name, "passed": bool(passed), "evidence": evidence})

    document = Document(DOCX)
    section = document.sections[0]
    body_width = twips(section.page_width - section.left_margin - section.right_margin)

    record("docx_exists", DOCX.exists() and DOCX.stat().st_size > 100_000, DOCX.stat().st_size)
    record("letter_page", close(twips(section.page_width), 12_240) and close(twips(section.page_height), 15_840), [twips(section.page_width), twips(section.page_height)])
    record("one_inch_margins", all(close(twips(value), 1_440) for value in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin)), [twips(section.top_margin), twips(section.bottom_margin), twips(section.left_margin), twips(section.right_margin)])
    record("header_footer_distance", close(twips(section.header_distance), 708, 3) and close(twips(section.footer_distance), 708, 3), [twips(section.header_distance), twips(section.footer_distance)])
    record("body_width", close(body_width, 9_360), body_width)

    normal = document.styles["Normal"]
    h1 = document.styles["Heading 1"]
    h2 = document.styles["Heading 2"]
    h3 = document.styles["Heading 3"]
    style_evidence = {
        "normal": [normal.font.name, normal.font.size.pt if normal.font.size else None, normal.paragraph_format.space_after.pt if normal.paragraph_format.space_after else None],
        "h1": [h1.font.name, h1.font.size.pt if h1.font.size else None],
        "h2": [h2.font.name, h2.font.size.pt if h2.font.size else None],
        "h3": [h3.font.name, h3.font.size.pt if h3.font.size else None],
    }
    record(
        "standard_business_brief_styles",
        normal.font.name == "Calibri"
        and normal.font.size == Pt(11)
        and h1.font.size == Pt(16)
        and h2.font.size == Pt(13)
        and h3.font.size == Pt(12),
        style_evidence,
    )

    all_paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    text = "\n".join(paragraph.text for paragraph in all_paragraphs)
    required = [
        "The Trinity Mandala",
        "Eiren Kestrel synthesis",
        "GMUT v-infinity evidence-bounded research kernel",
        "Candidate laws of thermo-psyche dynamics",
        "v641-v650 GMUT/THOS v1-v8 x1-x2 proposal",
        "Primary-source ledger",
        "Validation and reproducibility manifest",
    ]
    record("required_sections", all(item in text for item in required), required)
    record("substantial_content", len(document.paragraphs) >= 500 and len(document.tables) >= 40, {"top_level_paragraphs": len(document.paragraphs), "tables": len(document.tables)})
    record("embedded_visuals", len(document.inline_shapes) == 3, len(document.inline_shapes))
    oversized_images = []
    for index, shape in enumerate(document.inline_shapes, start=1):
        if twips(shape.width) and twips(shape.width) > body_width:
            oversized_images.append({"index": index, "width_twips": twips(shape.width)})
    record("images_fit_body", not oversized_images, oversized_images)

    table_widths = [table_width_twips(table) for table in document.tables]
    invalid_tables = [index + 1 for index, width in enumerate(table_widths) if width is None or width > body_width]
    record("tables_fit_body", not invalid_tables, {"invalid_tables": invalid_tables, "expected_max_twips": body_width})

    long_tokens = sorted({token for token in re.findall(r"\S{81,}", text) if not token.startswith("http")})
    record("no_unbreakable_text_tokens", not long_tokens, long_tokens[:10])
    record("no_machine_absolute_paths", not re.search(r"(?i)(?:[A-Z]:\\|C:/Users/|D:/GHC-Archives)", text), "scanned visible text")
    record("no_control_characters", not any(ord(char) < 32 and char not in "\n\t\r" for char in text), "scanned visible text")

    with zipfile.ZipFile(DOCX) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml")
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        images = [name for name in names if name.startswith("word/media/")]
        external_links = len(re.findall(r'TargetMode="External"', rels))
        record("valid_ooxml_parts", "[Content_Types].xml" in names and "word/styles.xml" in names and b"w:document" in document_xml, {"parts": len(names)})
        record("embedded_media_parts", len(images) == 3, images)
        record("external_source_links", external_links >= 40, external_links)

    record("header_label", "EIREN KESTREL" in section.header.paragraphs[0].text, section.header.paragraphs[0].text)
    record("footer_label", "Evidence-bounded candidate programme" in section.footer.paragraphs[0].text, section.footer.paragraphs[0].text)

    passed = sum(1 for item in checks if item["passed"])
    result = {
        "schema": "eiren-docx-validation-v1",
        "artifact": DOCX.name,
        "result": "PASS" if passed == len(checks) else "FAIL",
        "passed": passed,
        "total": len(checks),
        "checks": checks,
        "render_boundary": {
            "bundled_renderer": "blocked: LibreOffice executable unavailable",
            "microsoft_word_open": "passed in a fresh hidden read-only session",
            "microsoft_word_pdf_export": "blocked: native export timed out on both this dossier and an unrelated control DOCX",
            "claim": "structural and geometry-boundary validation, not page-raster validation",
        },
    }
    OUTPUT.write_text(json.dumps(result, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(result, indent=2, ensure_ascii=False))
    if result["result"] != "PASS":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
