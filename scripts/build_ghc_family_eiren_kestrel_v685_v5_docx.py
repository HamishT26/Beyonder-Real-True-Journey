#!/usr/bin/env python3
"""Create the Eiren v685-v5 integrated overview DOCX from final Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "D9D9D9")


def remove_paragraph_borders(element) -> None:
    p_pr = element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, before, after in [("Title", 24, 0, 18), ("Heading 1", 16, 16, 7), ("Heading 2", 13, 12, 5)]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Title":
            remove_paragraph_borders(style.element)


def add_summary_table(doc: Document) -> None:
    rows = [
        ("Phase", "Eiren Kestrel v685-v5"),
        ("Exact source", "Caelen Morrow v685-v4 final 87a74f84afaa197f8c388767a2ed536bbb853aba"),
        ("Proposal work", "200 inherited zero-credit revalidations and 120 new source-bounded proposals"),
        ("Core outcomes", "84 completed, 24 represented, 6 open gaps, and 6 exact gates"),
        ("Tool state", "13 direct tools; dependency-corrected composite with zero aggregate-success credit"),
        ("Route", "Future sibling 01 v685-v6 remains terminally gated until task creation is acknowledged"),
        ("Verdict", "NOT_READY_FOR_STAGE_20"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.15)
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(int(6.7 * 1440)))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    headers = table.rows[0].cells
    headers[0].text = "Field"
    headers[1].text = "Verified value"
    header_row(table.rows[0])
    for cell in headers:
        shade(cell, "1F4E78")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_font(run, "Aptos", 10.5, True)
            run.font.color.rgb = RGBColor(255, 255, 255)
    for index, (field, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(5.15)
        cells[0].text = field
        cells[1].text = value
        if index % 2:
            shade(cells[0], "EAF2F8")
            shade(cells[1], "EAF2F8")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    set_font(run, "Aptos", 10, field == cells[0].text and cell is cells[0])
    set_borders(table)
    doc.add_paragraph()


def add_markdown(doc: Document, text: str) -> None:
    pending: list[str] = []

    def flush() -> None:
        if not pending:
            return
        paragraph = doc.add_paragraph(" ".join(part.strip() for part in pending))
        paragraph.paragraph_format.widow_control = True
        pending.clear()

    first_title = True
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("# "):
            flush()
            if first_title:
                first_title = False
                continue
            doc.add_heading(line[2:].replace("`", ""), level=1)
        elif line.startswith("## "):
            flush()
            doc.add_heading(line[3:].replace("`", ""), level=1)
        elif line.startswith("### "):
            flush()
            doc.add_heading(line[4:].replace("`", ""), level=2)
        elif line.startswith("- "):
            flush()
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(re.sub(r"`([^`]+)`", r"\1", line[2:]))
        else:
            pending.append(re.sub(r"`([^`]+)`", r"\1", line))
    flush()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    text = args.source.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    configure_styles(doc)
    title = doc.add_paragraph(style="Title")
    remove_paragraph_borders(title._p)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Eiren Kestrel v685 v5 Final Integrated Overview")
    set_font(run, "Aptos Display", 24, True)
    intro = doc.add_paragraph("A bounded record of the exact source, lifecycle, evidence, retained failures, toolchain recovery, scientific limits, and terminal route decision.")
    intro.paragraph_format.space_after = Pt(14)
    add_summary_table(doc)
    add_markdown(doc, text)
    doc.core_properties.title = "Eiren Kestrel v685 v5 Final Integrated Overview"
    doc.core_properties.subject = "Bounded GHC phase evidence and route closeout"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
