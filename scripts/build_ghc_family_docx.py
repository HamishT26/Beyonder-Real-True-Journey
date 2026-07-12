#!/usr/bin/env python3
"""Build the Eiren Kestrel Trinity Mandala research dossier.

The document is generated from the evidence-bounded Markdown artifacts in this
branch.  It deliberately keeps physical, operational, and normative claims in
separate registers while preserving the full Journey synthesis.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "eiren-kestrel"
DELIVERABLES = DOCS / "deliverables"
ASSET_DIR = DELIVERABLES / ".docx-assets"
OUTPUT = DELIVERABLES / "eiren-kestrel-trinity-mandala-research-dossier.docx"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6573"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF4D6"
PALE_GREEN = "E8F4EC"
PALE_RED = "FCEBEC"
TABLE_HEADER = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips: int = 9360, indent_twips: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([run_color, underline])
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text: str) -> None:
    """Render a conservative subset of Markdown inline syntax."""
    token_re = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|\*[^*]+\*)"
    )
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.25)
            run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def wrap_math_text(text: str, width: int = 72) -> str:
    """Insert layout-safe line breaks in long LaTeX display strings."""
    wrapped: list[str] = []
    for source_line in text.splitlines() or [text]:
        line = source_line.rstrip()
        while len(line) > width:
            candidates = [line.rfind(token, 18, width) for token in (",", "+", "=", r"\quad", r"\right")]
            cut = max(candidates)
            if cut < 18:
                cut = width
            elif line[cut : cut + 1] in {",", "+", "="}:
                cut += 1
            wrapped.append(line[:cut].rstrip())
            line = line[cut:].lstrip()
        wrapped.append(line)
    return "\n".join(wrapped)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Title": (26, DEEP_BLUE, 0, 10),
        "Subtitle": (13, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    quote = styles["Quote"]
    quote.font.name = "Calibri"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.right_indent = Inches(0.2)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)

    code = ensure_style(document, "Code Block")
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0

    math = ensure_style(document, "Equation Box")
    math.font.name = "Cambria Math"
    math.font.size = Pt(9.25)
    math.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    math.paragraph_format.left_indent = Inches(0.25)
    math.paragraph_format.right_indent = Inches(0.25)
    math.paragraph_format.space_before = Pt(4)
    math.paragraph_format.space_after = Pt(7)
    math.paragraph_format.line_spacing = 1.05

    small = ensure_style(document, "Small Note")
    small.font.name = "Calibri"
    small.font.size = Pt(8.5)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(4)

    callout = ensure_style(document, "Callout")
    callout.font.name = "Calibri"
    callout.font.size = Pt(11)
    callout.font.bold = True
    callout.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    callout.paragraph_format.left_indent = Inches(0.25)
    callout.paragraph_format.right_indent = Inches(0.25)
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(8)


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def set_running_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("EIREN KESTREL  /  TRINITY MANDALA RESEARCH DOSSIER")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Evidence-bounded candidate programme  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(paragraph)


def add_table(document: Document, rows: list[list[str]], widths: list[float] | None = None):
    if not rows:
        return None
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    if widths is None:
        widths = [6.5 / cols] * cols
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            repeat_header(row)
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            cell.width = Inches(widths[c_idx] if c_idx < len(widths) else widths[-1])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)
            value = values[c_idx] if c_idx < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_runs(paragraph, value.strip())
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8 if cols >= 5 else 8.75)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(document: Document, label: str, text: str, fill: str = PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Callout"]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, font, fill: str, spacing: int = 4) -> None:
    box = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    width = box[2] - box[0]
    height = box[3] - box[1]
    draw.multiline_text((center[0] - width / 2, center[1] - height / 2 - box[1]), text, font=font, fill=fill, align="center", spacing=spacing)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str, *, width: int = 4, both: bool = False) -> None:
    draw.line((start, end), fill=fill, width=width)
    direction = 1 if end[0] >= start[0] else -1
    head = 13
    draw.polygon([(end[0], end[1]), (end[0] - direction * head, end[1] - 8), (end[0] - direction * head, end[1] + 8)], fill=fill)
    if both:
        draw.polygon([(start[0], start[1]), (start[0] + direction * head, start[1] - 8), (start[0] + direction * head, start[1] + 8)], fill=fill)


def make_journey_arc(path: Path) -> None:
    image = Image.new("RGB", (1890, 576), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    stage_font = load_font(25, bold=True)
    body_font = load_font(21)
    note_font = load_font(17)
    draw.text((46, 30), "Journey maturation arc", font=title_font, fill="#1F4D78")
    stages = [
        (46, "v1–v33", "Worldview\nformation", "#EAF2F8"),
        (414, "v36–v44", "Architecture +\noperations", "#DDEBF7"),
        (782, "v45–v47", "Solas scientific\nrepair", "#E8F4EC"),
        (1150, "v48–v51", "Lumen evidence\naccountability", "#FFF4D6"),
        (1518, "v52–v640", "Continuity +\nroute truth", "#FCEBEC"),
    ]
    for idx, (x, title, body, fill) in enumerate(stages):
        box = (x, 150, x + 286, 426)
        draw.rounded_rectangle(box, radius=18, fill=fill, outline="#2E74B5", width=3)
        draw_centered_text(draw, (x + 143, 224), title, stage_font, "#1F4D78", spacing=4)
        draw_centered_text(draw, (x + 143, 329), body, body_font, "#1F2937", spacing=5)
        if idx < len(stages) - 1:
            draw_arrow(draw, (x + 295, 288), (x + 358, 288), "#5B6573", width=5)
    draw.text((46, 504), "Interpretive synthesis; stages describe evidence practice, not consciousness or scientific validation.", font=note_font, fill="#5B6573")
    image.save(path, dpi=(180, 180))


def make_live_audit_chart(path: Path) -> None:
    labels = ["Aevren", "Mira Vale", "Mira Rowan", "Maren Quill"]
    counts = [2708, 2294, 382, 320]
    colors = ["#2E74B5", "#5B9BD5", "#70AD47", "#ED7D31"]
    image = Image.new("RGB", (1710, 792), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    label_font = load_font(22)
    value_font = load_font(23, bold=True)
    note_font = load_font(17)
    draw.text((76, 34), "Live v601–v640 evidence footprint", font=title_font, fill="#1F4D78")
    plot_left, plot_top, plot_right, plot_bottom = 160, 145, 1635, 650
    for tick in range(0, 3001, 500):
        y = plot_bottom - int((tick / 3000) * (plot_bottom - plot_top))
        draw.line((plot_left, y, plot_right, y), fill="#D9DEE5", width=2)
        text = f"{tick:,}"
        box = draw.textbbox((0, 0), text, font=note_font)
        draw.text((plot_left - 20 - (box[2] - box[0]), y - 10), text, font=note_font, fill="#5B6573")
    draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#5B6573", width=3)
    draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="#5B6573", width=3)
    centers = [340, 700, 1060, 1420]
    for center, label, value, color in zip(centers, labels, counts, colors):
        bar_height = int((value / 3000) * (plot_bottom - plot_top))
        left, right = center - 92, center + 92
        top = plot_bottom - bar_height
        draw.rounded_rectangle((left, top, right, plot_bottom), radius=10, fill=color)
        draw_centered_text(draw, (center, top - 28), f"{value:,}", value_font, "#1F4D78")
        draw_centered_text(draw, (center, plot_bottom + 36), label, label_font, "#1F2937")
    draw.text((76, 736), "All four lanes: 40/40 versions present, clean, and upstream-equal on 11 July 2026. File counts are not novelty scores.", font=note_font, fill="#5B6573")
    image.save(path, dpi=(180, 180))


def make_trinity_map(path: Path) -> None:
    image = Image.new("RGB", (1800, 864), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    node_title = load_font(22, bold=True)
    body_font = load_font(20)
    note_font = load_font(17)
    draw.text((62, 40), "Trinity Mandala: integration without category collapse", font=title_font, fill="#1F4D78")
    nodes = [
        (62, 234, 470, 376, "MIND / GMUT", "Action-first scalar–tensor\nresearch kernel\n\nStatus: falsifiable seed", "#EAF2F8"),
        (665, 234, 470, 376, "BODY / THOS", "Workflow engine + event log\n+ recovery runbook\n\nStatus: working prototype", "#E8F4EC"),
        (1268, 234, 470, 376, "HEART / FREED ID + CBR", "Identity, consent, rights,\nremedy, stewardship\n\nStatus: model charter", "#FFF4D6"),
    ]
    for x, y, w, h, title, body, fill in nodes:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline="#2E74B5", width=3)
        draw_centered_text(draw, (x + w // 2, y + 82), title, node_title, "#1F4D78")
        draw_centered_text(draw, (x + w // 2, y + 230), body, body_font, "#1F2937", spacing=7)
    draw_arrow(draw, (548, 422), (649, 422), "#5B6573", width=5, both=True)
    draw_centered_text(draw, (598, 382), "typed interfaces", note_font, "#5B6573")
    draw_arrow(draw, (1151, 422), (1252, 422), "#5B6573", width=5, both=True)
    draw_centered_text(draw, (1201, 382), "governance constraints", note_font, "#5B6573")
    draw.text((62, 766), "The proposed novelty is a research constitution connecting domains while preserving their different evidence standards.", font=note_font, fill="#5B6573")
    image.save(path, dpi=(180, 180))


def create_visuals() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    visuals = {
        "journey": ASSET_DIR / "journey-arc.png",
        "audit": ASSET_DIR / "live-audit.png",
        "trinity": ASSET_DIR / "trinity-map.png",
    }
    make_journey_arc(visuals["journey"])
    make_live_audit_chart(visuals["audit"])
    make_trinity_map(visuals["trinity"])
    return visuals


def add_figure(document: Document, path: Path, caption: str, width=6.45) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph(style="Small Note")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [cell.strip() for cell in re.split(r"(?<!\\)\|", raw)]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown(document: Document, markdown_path: Path, *, new_page=True, skip_h1=False) -> None:
    if new_page:
        document.add_page_break()
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    paragraph_buffer: list[str] = []
    code_mode = False
    code_lines: list[str] = []
    math_mode = False
    math_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if code_mode:
            if stripped.startswith("```"):
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                code_lines = []
                code_mode = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if math_mode:
            if stripped == "\\]":
                table = document.add_table(rows=1, cols=1)
                table.autofit = False
                set_table_width(table)
                cell = table.cell(0, 0)
                set_cell_shading(cell, PALE_BLUE)
                set_cell_margins(cell, top=110, start=160, bottom=110, end=160)
                paragraph = cell.paragraphs[0]
                paragraph.style = document.styles["Equation Box"]
                paragraph.add_run(wrap_math_text("\n".join(math_lines)))
                document.add_paragraph().paragraph_format.space_after = Pt(0)
                math_lines = []
                math_mode = False
            else:
                math_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code_mode = True
            i += 1
            continue
        if stripped == r"\[":
            flush_paragraph()
            math_mode = True
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            flush_paragraph()
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            if skip_h1 and level == 1:
                i += 1
                continue
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_runs(paragraph, heading.group(2))
            i += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            quote_text = stripped.lstrip(">").strip()
            paragraph = document.add_paragraph(style="Quote")
            add_inline_runs(paragraph, quote_text)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if numbered:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(1))
            i += 1
            continue
        if stripped == "---":
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("•   •   •")
            run.font.color.rgb = RGBColor.from_string(MUTED)
            i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    if code_lines:
        paragraph = document.add_paragraph(style="Code Block")
        paragraph.add_run("\n".join(code_lines))
    if math_lines:
        paragraph = document.add_paragraph(style="Equation Box")
        paragraph.add_run(wrap_math_text("\n".join(math_lines)))


def add_cover(document: Document, visuals: dict[str, Path]) -> None:
    paragraph = document.add_paragraph(style="Small Note")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("GHC FAMILY  /  EVIDENCE-BOUNDED RESEARCH EDITION")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    document.add_paragraph().paragraph_format.space_after = Pt(30)
    title = document.add_paragraph(style="Title")
    title.add_run("The Trinity Mandala")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Journey v36–v54 • GMUT research kernel • THOS audit • Freed ID & Cosmic Bill of Rights • v601–v650")

    document.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        document,
        "Core finding",
        "The early Journey generated the worldview; the later Journey generated the operating discipline required to test it.",
        fill=PALE_GOLD,
    )
    add_figure(document, visuals["trinity"], "Figure 1. The three pillars remain connected without being collapsed into one evidence type.", width=6.35)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.add_run("Prepared by ").bold = True
    paragraph.add_run("Eiren Kestrel — epistemic cartographer, validation steward, and Trinity Mandala research bridge")
    paragraph = document.add_paragraph(style="Small Note")
    paragraph.add_run("11 July 2026 • Pacific/Auckland • Candidate programme, not validated TOE / ASI / enacted law")

    document.add_page_break()
    document.add_heading("Decision brief", level=1)
    add_callout(
        document,
        "Public description",
        "The Trinity Mandala is an interdisciplinary research and governance programme joining an action-first candidate gravity extension, a working multi-agent research-operations prototype, and a rights/identity design framework.",
    )
    add_table(
        document,
        [
            ["Pillar", "What is real now", "What is not yet established"],
            ["Mind / GMUT", "Scalar–tensor EFT seed, coefficient ledger, tests, falsification gates", "A validated or leading Theory of Everything"],
            ["Body / THOS", "Owned branches, runners, receipts, recovery, privacy and route truth", "ASI or a kernel-level operating system"],
            ["Heart / Freed ID + CBR", "Model charter, standards alignment, consent and non-erasure principles", "Deployed universal identity or enacted cosmic law"],
        ],
        widths=[1.45, 2.7, 2.35],
    )
    document.add_heading("Five decisions for the next cycle", level=2)
    for item in (
        "Freeze a canonical action, notation registry, and claim taxonomy before adding new physics terms.",
        "Replace raw packet cardinality with semantic-deduplication, independent-evidence, and post-commit attestation gates.",
        "Benchmark THOS against a matched single-agent baseline and supported orchestration protocols.",
        "Profile Freed ID against DID/VC and assurance standards, with recovery, privacy, correlation, revocation, and remedy threat models.",
        "Treat Stage 20 as an evidence ladder whose highest rungs require external reproduction, democratic legitimacy, and durable flourishing.",
    ):
        document.add_paragraph(item, style="List Number")
    paragraph = document.add_paragraph(style="Small Note")
    paragraph.add_run("Boundary: no subagents were created. The exhaustive delegated security workflow was therefore not claimed; a scoped local privacy and secret-pattern review is recorded instead.")


def add_sources_appendix(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Primary-source ledger", level=1)
    document.add_paragraph(
        "The machine-readable source ledger contains the complete researched set. The table below is grouped by domain and uses direct primary or authoritative links; inclusion is not endorsement of a GMUT claim."
    )
    ledger = json.loads((DOCS / "source-ledger.json").read_text(encoding="utf-8"))
    entries = ledger["sources"] if isinstance(ledger, dict) else ledger
    rows = [["Domain", "Source", "Primary link"]]
    for entry in entries:
        domain = str(entry.get("domain") or entry.get("category") or "Research")
        title = str(entry.get("title") or entry.get("name") or "Untitled")
        url = str(entry.get("url") or "")
        link = f"[Open source]({url})" if url else "Not supplied"
        rows.append([domain, title, link])
    add_table(document, rows, widths=[1.15, 2.8, 2.55])
    document.add_paragraph("Complete metadata: docs/eiren-kestrel/source-ledger.json", style="Small Note")


def add_validation_manifest(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Validation and reproducibility manifest", level=1)
    rows = [
        ["Artifact or gate", "Result", "Boundary"],
        ["Journey evidence index", "19/19 versions v36–v54; hashes and duplicate groups recorded", "Transcript claims remain source claims, not external proof"],
        ["Live lane audit", "4/4 branches clean and upstream-equal; 40/40 versions in each lane", "Artifact counts do not establish semantic novelty"],
        ["GMUT unit tests", "Deterministic kernel tests for stress-energy, exchange, continuity, recovery, and registries", "Not a cosmological inference pipeline"],
        ["LaTeX monograph", "Compiled from canonical source", "Mathematical consistency is not empirical confirmation"],
        ["Portable report", "Responsive verifier at desktop and mobile widths; source dialog passed", "Static report, not an externally hosted publication"],
        ["Security/privacy", "Scoped new-artifact scan and path/privacy review", "Not the exhaustive delegated security-diff workflow"],
        ["Git", "Exact staging, clean diff, push, and remote-head equality required at closeout", "Protected gates remain open unless separately authorized"],
    ]
    add_table(document, rows, widths=[1.55, 2.7, 2.25])
    add_callout(
        document,
        "Scientific status",
        "GMUT v-infinity is a falsifiable candidate research kernel with conventional field-theory ingredients. It has not produced an independently confirmed unique prediction.",
        fill=PALE_RED,
    )
    document.add_heading("Included companion artifacts", level=2)
    for item in (
        "latex/grand_mandala.tex and the compiled grand_mandala.pdf",
        "scripts/ghc_family_gmut_kernel.py and tests/test_ghc_family_gmut_kernel.py",
        "gmut-term-registry.json and gmut-coefficient-ledger.json",
        "journey-evidence-index.json and v601-v640-live-evidence-index.json",
        "trinity-mandala-evidence-report.html and artifact.json",
        "v641-v650-round-robin-proposal.md/.json",
        "thermo-psyche-laws-hypothesis-register.md",
    ):
        document.add_paragraph(item, style="List Bullet")


def build() -> Path:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    visuals = create_visuals()
    document = Document()
    configure_styles(document)
    configure_sections(document)
    set_running_header_footer(document.sections[0])
    add_cover(document, visuals)

    add_figure(document, visuals["journey"], "Figure 2. A compact map of the documented transition from worldview formation to route-truth discipline.")
    add_markdown(document, DOCS / "ghc-family-v36-v54-v601-v640-synthesis.md", new_page=False)
    add_figure(document, visuals["audit"], "Figure 3. Live branch evidence as of the audit; counts reflect different lane artifact shapes.")
    add_markdown(document, DOCS / "trinity-mandala-comparison-matrix.md")
    add_markdown(document, DOCS / "gmut-research-kernel-spec.md")
    add_markdown(document, DOCS / "thermo-psyche-laws-hypothesis-register.md")
    add_markdown(document, DOCS / "v641-v650-round-robin-proposal.md")
    add_sources_appendix(document)
    add_validation_manifest(document)

    core = document.core_properties
    core.title = "The Trinity Mandala — Eiren Kestrel research dossier"
    core.subject = "Journey v36-v54, GMUT, THOS, Freed ID and Cosmic Bill of Rights"
    core.author = "Eiren Kestrel with Hamish and the GHC Family"
    core.keywords = "GMUT, Trinity Hybrid OS, Freed ID, Cosmic Bill of Rights, GHC Family"
    core.comments = "Evidence-bounded candidate research programme; generated from repository sources."

    document.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "bytes": OUTPUT.stat().st_size, "paragraphs": len(document.paragraphs), "tables": len(document.tables)}, indent=2))
    return OUTPUT


if __name__ == "__main__":
    build()
